"""
StudySync Backend  v5
=====================
Changes from v4:
  - Topic-wise section hataya (/analyze, /merge removed)
  - Smart merge ka duplicate detection fix kiya:
      * Exact match pehle hatao (normalized text)
      * TF-IDF threshold 0.45 → 0.65 (zyada strict)
      * Subword overlap check add kiya
  - Sirf smart PDF section raha

Features:
  POST /upload            - Files upload
  GET  /files             - File list
  DELETE /files/<name>    - File delete
  POST /smart-merge       - Smart line merge (MAIN feature)
  POST /summary           - Auto summary
  POST /keywords          - Top keywords
  POST /important-lines   - Important sentences
  POST /search            - Text search
  POST /flashcards        - Q&A flashcards
  POST /quiz              - MCQ quiz
  POST /stats             - File statistics
  POST /compare           - 2 files compare
  POST /export-docx       - Word export
  GET  /download/<name>   - Download file
  GET  /outputs           - Output files list
  POST /clear             - Clear all
  GET  /health            - Server health

Install:
  pip install flask flask-cors pdfplumber reportlab scikit-learn python-docx

Run:
  python backend.py  →  http://localhost:5000
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os, re
from collections import Counter
from datetime import datetime

# ── Optional imports ──────────────────────────────────────────────────────────
try:
    import pdfplumber;  PDF_OK = True
except ImportError:
    PDF_OK = False;     print("[WARN] pip install pdfplumber")

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    ML_OK = True
except ImportError:
    ML_OK = False;      print("[WARN] pip install scikit-learn")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle)
    RL_OK = True
except ImportError:
    RL_OK = False;      print("[WARN] pip install reportlab")

try:
    import docx as _docx; DOCX_OK = True
except ImportError:
    DOCX_OK = False;    print("[WARN] pip install python-docx")

# ══════════════════════════════════════════════════════════════════════════════
app = Flask(__name__)
CORS(app)

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Thresholds ────────────────────────────────────────────────────────────────
TFIDF_THRESHOLD  = 0.65   # Higher = zyada strict dedup (was 0.45 — too loose)
WORD_OVERLAP_MIN = 0.70   # 70%+ common words = duplicate

# ── Junk patterns ─────────────────────────────────────────────────────────────
JUNK_RE = re.compile(
    r'^\s*$'
    r'|^\s*\d+\s*$'
    r'|^page\s*\d+'
    r'|^\s*.{1,4}\s*$'
    r'|www\.|http|https|@'
    r'|^(copyright|all rights|©)'
    r'|^(figure|fig\.|table|chart)\s*\d+'
    r'|^[\-_=*]{3,}$'
    r'|^\s*[•\-–]\s*$',
    re.IGNORECASE
)

STOP_WORDS = {
    "the","is","in","it","of","and","a","to","was","that","he","she",
    "they","we","you","i","this","with","for","on","are","be","as",
    "at","have","from","or","an","but","not","what","all","were",
    "when","there","been","one","had","by","do","so","up","which",
    "their","will","about","if","would","no","said","its","into","has"
}

# ══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_text(filepath: str) -> str:
    ext = filepath.rsplit(".", 1)[-1].lower()
    try:
        if ext == "pdf":    return _pdf_text(filepath)
        elif ext == "docx": return _docx_text(filepath)
        else:               return _txt_text(filepath)
    except Exception as e:
        return f"[Error: {e}]"

def _pdf_text(path):
    if not PDF_OK: return "[pdfplumber install karo]"
    pages = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            t = p.extract_text()
            if t: pages.append(t.strip())
    return "\n\n".join(pages)

def _docx_text(path):
    if not DOCX_OK: return "[python-docx install karo]"
    doc = _docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def _txt_text(path):
    with open(path, encoding="utf-8", errors="ignore") as f:
        return f.read()

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x20-\x7E\n]', '', text)
    return text.strip()

def split_sentences(text: str) -> list:
    parts = re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in parts if len(s.strip()) > 20]

def split_lines(text: str) -> list:
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if len(line) > 200:
            lines.extend(
                s.strip() for s in re.split(r'(?<=[.!?])\s+', line) if s.strip()
            )
        else:
            lines.append(line)
    return lines

def _normalize(text: str) -> str:
    """Comparison ke liye text normalize karta hai."""
    t = text.lower().strip()
    t = re.sub(r'[^\w\s]', '', t)       # punctuation hatao
    t = re.sub(r'\s+', ' ', t)           # extra spaces hatao
    return t

def _load_texts(filenames: list) -> dict:
    return {n: extract_text(os.path.join(UPLOAD_DIR, n))
            for n in filenames
            if os.path.exists(os.path.join(UPLOAD_DIR, n))}

# ══════════════════════════════════════════════════════════════════════════════
# DEDUPLICATION (Fixed)
# ══════════════════════════════════════════════════════════════════════════════

def _exact_dedup(lines: list) -> list:
    """
    Step 1: Exact duplicates hatao (normalized comparison).
    Same sentence different punctuation/spacing ke saath bhi catch karta hai.
    """
    seen = set()
    result = []
    for line in lines:
        key = _normalize(line)
        if key not in seen:
            seen.add(key)
            result.append(line)
    return result

def _word_overlap_dedup(lines: list) -> list:
    """
    Step 2: Word overlap se near-duplicates hatao.
    Agar 2 lines ke 70%+ words same hain → duplicate.
    ML nahi chahiye — fast aur reliable.
    """
    def word_set(text):
        words = set(re.findall(r'[a-zA-Z]{3,}', text.lower()))
        return words - STOP_WORDS

    result = []
    for line in lines:
        ws = word_set(line)
        if not ws:
            result.append(line)
            continue
        is_dup = False
        for kept in result:
            kws = word_set(kept)
            if not kws: continue
            common = ws & kws
            smaller = min(len(ws), len(kws))
            if smaller > 0 and len(common) / smaller >= WORD_OVERLAP_MIN:
                is_dup = True
                break
        if not is_dup:
            result.append(line)
    return result

def _tfidf_dedup(lines: list) -> list:
    """
    Step 3: TF-IDF cosine similarity se remaining duplicates hatao.
    Threshold 0.65 — strict enough to catch paraphrases.
    """
    if len(lines) <= 1 or not ML_OK:
        return lines
    try:
        vec = TfidfVectorizer(stop_words="english", min_df=1, ngram_range=(1, 2))
        mat = vec.fit_transform(lines)
        sim = cosine_similarity(mat)
        kept, removed = [], set()
        for i in range(len(lines)):
            if i in removed: continue
            kept.append(lines[i])
            for j in range(i + 1, len(lines)):
                if j not in removed and sim[i][j] >= TFIDF_THRESHOLD:
                    removed.add(j)
        return kept
    except:
        return lines

def deduplicate(lines: list) -> list:
    """
    3-step deduplication pipeline:
    1. Exact match (normalized)
    2. Word overlap >= 70%
    3. TF-IDF cosine >= 0.65
    """
    step1 = _exact_dedup(lines)
    step2 = _word_overlap_dedup(step1)
    step3 = _tfidf_dedup(step2)
    print(f"[Dedup] {len(lines)} → {len(step1)} → {len(step2)} → {len(step3)} lines")
    return step3

# ══════════════════════════════════════════════════════════════════════════════
# SMART MERGE (Main Feature)
# ══════════════════════════════════════════════════════════════════════════════

def is_junk(line: str) -> bool:
    return bool(JUNK_RE.search(line.strip()))

def is_useful(line: str) -> bool:
    if len(line) < 15: return False
    return len(re.findall(r'[a-zA-Z]{2,}', line)) >= 2

def smart_merge(texts: dict) -> list:
    """
    Saari files ki lines leke:
    1. Split into lines
    2. Junk hatao
    3. Too short / no words hatao
    4. 3-step deduplication karo
    """
    all_lines = []
    for text in texts.values():
        all_lines.extend(split_lines(clean_text(text)))

    # Step 1: Junk + useless filter
    filtered = [l for l in all_lines if not is_junk(l) and is_useful(l)]

    if not filtered:
        return []

    # Step 2: 3-step dedup
    return deduplicate(filtered)

# ══════════════════════════════════════════════════════════════════════════════
# OTHER FEATURE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _score_sentences(sentences: list) -> list:
    if not ML_OK or len(sentences) < 2:
        return [(s, 1.0) for s in sentences]
    try:
        vec = TfidfVectorizer(stop_words="english")
        mat = vec.fit_transform(sentences)
        scores = mat.sum(axis=1).A1.tolist()
        max_s = max(scores) or 1
        return [(s, round(sc / max_s, 3)) for s, sc in zip(sentences, scores)]
    except:
        return [(s, 1.0) for s in sentences]

def _make_summary(text: str, max_sentences=5) -> str:
    sentences = split_sentences(clean_text(text))
    if not sentences: return ""
    scored = _score_sentences(sentences)
    scored.sort(key=lambda x: x[1], reverse=True)
    return " ".join(s for s, _ in scored[:max_sentences])

def _extract_keywords(texts_dict: dict, top_n=15) -> list:
    all_text = " ".join(texts_dict.values())
    words = re.findall(r'[a-zA-Z]{3,}', all_text.lower())
    filtered = [w for w in words if w not in STOP_WORDS]
    freq = Counter(filtered)
    return [{"word": w, "count": c} for w, c in freq.most_common(top_n)]

def _make_flashcards(text: str, count=10) -> list:
    sentences = split_sentences(clean_text(text))
    cards = []
    patterns = [
        r'^(.{3,40})\s+is\s+(.{10,})',
        r'^(.{3,40})\s+are\s+(.{10,})',
        r'^(.{3,40})\s+means\s+(.{10,})',
        r'^(.{3,40})\s+refers to\s+(.{10,})',
        r'^(.{3,40})\s+=\s+(.{10,})',
    ]
    for sent in sentences:
        for pat in patterns:
            m = re.match(pat, sent, re.IGNORECASE)
            if m:
                subj = m.group(1).strip()
                if len(subj.split()) <= 6:
                    cards.append({"question": f"What is {subj}?", "answer": sent})
                    break
        if len(cards) >= count: break

    if len(cards) < count:
        scored = _score_sentences(sentences)
        scored.sort(key=lambda x: x[1], reverse=True)
        for sent, sc in scored:
            if sc > 0.5 and len(cards) < count:
                if not any(c["answer"] == sent for c in cards):
                    words = sent.split()
                    if len(words) > 4:
                        cards.append({
                            "question": f"Explain: '{' '.join(words[:4])}...'",
                            "answer": sent,
                        })
    return cards[:count]

def _make_quiz(text: str, count=5) -> list:
    import random
    sentences = split_sentences(clean_text(text))
    if not sentences: return []
    scored = _score_sentences(sentences)
    scored.sort(key=lambda x: x[1], reverse=True)
    top = [s for s, sc in scored[:count * 3] if sc > 0.3]
    quiz = []
    for sent in top[:count]:
        words = sent.split()
        if len(words) < 6: continue
        candidates = [w for w in words if len(w) > 4 and w[0].isupper() and w != words[0]]
        if not candidates: candidates = [w for w in words if re.match(r'\d', w)]
        if not candidates: candidates = [w for w in words if len(w) > 5]
        if not candidates: continue
        answer = candidates[0]
        question_text = sent.replace(answer, "______", 1)
        others = [w for w in words if w != answer and len(w) > 3][:3]
        while len(others) < 3: others.append("None of these")
        options = [answer] + others[:3]
        random.shuffle(options)
        quiz.append({"question": question_text, "options": options, "answer": answer})
    return quiz

def _get_stats(text: str, filename: str) -> dict:
    clean = clean_text(text)
    words = re.findall(r'[a-zA-Z]+', clean)
    sentences = split_sentences(clean)
    return {
        "filename":         filename,
        "characters":       len(clean),
        "words":            len(words),
        "sentences":        len(sentences),
        "unique_words":     len(set(w.lower() for w in words)),
        "avg_word_length":  round(sum(len(w) for w in words) / len(words), 1) if words else 0,
        "avg_sent_len":     round(len(words) / len(sentences), 1) if sentences else 0,
        "estimated_pages":  max(1, round(len(words) / 250)),
        "read_time_min":    max(1, round(len(words) / 200)),
    }

# ══════════════════════════════════════════════════════════════════════════════
# PDF / TXT GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

PURPLE = "#5349C8"; LPURPLE = "#EEEDFE"; DPURPLE = "#3C3489"; GRAY = "#D3D1C7"

def _styles():
    b = getSampleStyleSheet()
    return {
        "title":    ParagraphStyle("T", parent=b["Title"], fontSize=20,
                        textColor=colors.HexColor(PURPLE), spaceAfter=4),
        "subtitle": ParagraphStyle("S", parent=b["Normal"], fontSize=9,
                        textColor=colors.grey, spaceAfter=16),
        "body":     ParagraphStyle("B", parent=b["Normal"], fontSize=10,
                        leading=16, spaceAfter=6),
        "source":   ParagraphStyle("Src", parent=b["Normal"], fontSize=8,
                        textColor=colors.grey, spaceAfter=3),
        "bullet":   ParagraphStyle("Bul", parent=b["Normal"], fontSize=10,
                        leading=17, leftIndent=12, spaceAfter=3),
    }

def smart_merge_pdf(lines: list, sources: list, outpath: str):
    if not RL_OK:
        return smart_merge_txt(lines, sources, outpath.replace(".pdf", ".txt"))
    doc = SimpleDocTemplate(outpath, pagesize=A4,
          leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    s = _styles(); story = []

    story.append(Paragraph("StudySync — Smart Merged Notes", s["title"]))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}  |  "
        f"{len(lines)} unique lines  |  "
        f"Sources: {', '.join(s2.replace('.pdf','').replace('.docx','') for s2 in sources)}",
        s["subtitle"]))

    info = Table([[Paragraph(
        "<b>Duplicates removed</b>: Yes (3-step)  |  "
        "<b>Junk removed</b>: Yes  |  "
        "<b>Topic grouping</b>: None",
        s["source"])]], colWidths=[16*cm])
    info.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(-1,-1), colors.HexColor(LPURPLE)),
        ("BOX",        (0,0),(-1,-1), 0.5, colors.HexColor("#AFA9EC")),
        ("TOPPADDING", (0,0),(-1,-1), 8),
        ("BOTTOMPADDING",(0,0),(-1,-1), 8),
        ("LEFTPADDING", (0,0),(-1,-1), 10),
    ]))
    story += [info, Spacer(1, 0.4*cm),
              HRFlowable(width="100%", thickness=1, color=colors.HexColor("#AFA9EC")),
              Spacer(1, 0.3*cm)]

    for line in lines:
        if len(line) < 80 and not line.endswith('.') and line[0].isupper():
            story += [Spacer(1, 0.15*cm), Paragraph(f"<b>{line}</b>", s["body"])]
        else:
            story.append(Paragraph(f"• {line}", s["bullet"]))

    doc.build(story)
    print(f"[PDF] {outpath}")

def smart_merge_txt(lines: list, sources: list, outpath: str):
    out = [
        "=" * 60,
        "  STUDYSYNC — SMART MERGED NOTES",
        f"  Generated : {datetime.now().strftime('%d %b %Y')}",
        f"  Sources   : {', '.join(sources)}",
        f"  Lines     : {len(lines)}",
        "  Dedup     : Exact + Word Overlap + TF-IDF",
        "=" * 60, "",
    ]
    out += [f"• {l}" for l in lines] + [""]
    open(outpath, "w", encoding="utf-8").write("\n".join(out))
    print(f"[TXT] {outpath}")

def export_docx(lines: list, sources: list, outpath: str) -> bool:
    if not DOCX_OK: return False
    doc = _docx.Document()
    doc.add_heading("StudySync — Smart Merged Notes", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}")
    doc.add_paragraph(f"Sources: {', '.join(sources)}")
    doc.add_paragraph(f"Total unique lines: {len(lines)}")
    doc.add_page_break()
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")
    doc.save(outpath)
    return True

# ══════════════════════════════════════════════════════════════════════════════
# FLASK ROUTES
# ══════════════════════════════════════════════════════════════════════════════

# ── Health ────────────────────────────────────────────────────────────────────
@app.route("/health")
def health():
    return jsonify({
        "status": "online", "version": "v5",
        "pdfplumber": PDF_OK, "sklearn": ML_OK,
        "reportlab": RL_OK, "python_docx": DOCX_OK,
        "tfidf_threshold": TFIDF_THRESHOLD,
        "word_overlap_min": WORD_OVERLAP_MIN,
        "time": datetime.now().strftime("%d %b %Y %I:%M %p"),
    })

# ── Upload ────────────────────────────────────────────────────────────────────
@app.route("/upload", methods=["POST"])
def upload():
    if "files" not in request.files:
        return jsonify({"error": "Koi file nahi mili"}), 400
    saved = []
    for f in request.files.getlist("files"):
        if f.filename.endswith((".pdf", ".docx", ".txt")):
            path = os.path.join(UPLOAD_DIR, f.filename)
            f.save(path)
            saved.append({"name": f.filename,
                           "size_kb": round(os.path.getsize(path) / 1024, 1)})
    return jsonify({"uploaded": saved, "count": len(saved)})

# ── Files list ────────────────────────────────────────────────────────────────
@app.route("/files")
def list_files():
    files = [{"name": f,
               "size_kb": round(os.path.getsize(os.path.join(UPLOAD_DIR, f)) / 1024, 1)}
             for f in os.listdir(UPLOAD_DIR)
             if os.path.isfile(os.path.join(UPLOAD_DIR, f))]
    return jsonify({"files": files, "count": len(files)})

# ── Delete file ───────────────────────────────────────────────────────────────
@app.route("/files/<filename>", methods=["DELETE"])
def delete_file(filename):
    path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(path):
        os.remove(path)
        return jsonify({"deleted": filename})
    return jsonify({"error": "File nahi mili"}), 404

# ── Smart Merge (Main) ────────────────────────────────────────────────────────
@app.route("/smart-merge", methods=["POST"])
def smart_merge_route():
    """
    Main feature — line-by-line smart merge with 3-step deduplication.
    Body: { "files": ["a.pdf", "b.pdf"], "format": "pdf" }
    """
    data    = request.json or {}
    texts   = _load_texts(data.get("files", []))
    fmt     = data.get("format", "pdf")

    if not texts:
        return jsonify({"error": "Koi valid file nahi mili"}), 400

    lines   = smart_merge(texts)
    sources = list(texts.keys())
    outpath = os.path.join(OUTPUT_DIR, f"smart_merged_notes.{fmt}")

    if fmt == "pdf":
        smart_merge_pdf(lines, sources, outpath)
    else:
        smart_merge_txt(lines, sources, outpath)

    return jsonify({
        "filename":     f"smart_merged_notes.{fmt}",
        "unique_lines": len(lines),
        "lines":        lines,
        "sources":      sources,
    })

# ── Summary ───────────────────────────────────────────────────────────────────
@app.route("/summary", methods=["POST"])
def summary():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    n     = data.get("sentences", 5)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400
    summaries = {fname: _make_summary(text, n) for fname, text in texts.items()}
    combined  = _make_summary(" ".join(texts.values()), n)
    return jsonify({"individual": summaries, "combined": combined})

# ── Keywords ──────────────────────────────────────────────────────────────────
@app.route("/keywords", methods=["POST"])
def keywords():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    top   = data.get("top", 15)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400
    per_file = {f: _extract_keywords({f: t}, top) for f, t in texts.items()}
    return jsonify({"per_file": per_file, "combined": _extract_keywords(texts, top)})

# ── Important lines ───────────────────────────────────────────────────────────
@app.route("/important-lines", methods=["POST"])
def important_lines():
    data      = request.json or {}
    texts     = _load_texts(data.get("files", []))
    top       = data.get("top", 20)
    min_score = data.get("min_score", 0.3)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400
    combined  = " ".join(texts.values())
    sentences = split_sentences(clean_text(combined))
    scored    = _score_sentences(sentences)
    scored.sort(key=lambda x: x[1], reverse=True)
    result    = [{"line": s, "score": sc} for s, sc in scored if sc >= min_score][:top]
    return jsonify({"lines": result, "total": len(result)})

# ── Search ────────────────────────────────────────────────────────────────────
@app.route("/search", methods=["POST"])
def search():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    query = data.get("query", "").lower().strip()
    if not query: return jsonify({"error": "Query deni hogi"}), 400
    results = {}
    for fname, text in texts.items():
        matches = [s for s in split_sentences(clean_text(text)) if query in s.lower()]
        if matches: results[fname] = {"matches": matches, "count": len(matches)}
    total = sum(r["count"] for r in results.values())
    return jsonify({"query": query, "results": results,
                    "total_matches": total, "files_matched": len(results)})

# ── Flashcards ────────────────────────────────────────────────────────────────
@app.route("/flashcards", methods=["POST"])
def flashcards():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    count = data.get("count", 10)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400
    cards = _make_flashcards(" ".join(texts.values()), count)
    return jsonify({"flashcards": cards, "total": len(cards)})

# ── Quiz ──────────────────────────────────────────────────────────────────────
@app.route("/quiz", methods=["POST"])
def quiz():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    count = data.get("count", 5)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400
    questions = _make_quiz(" ".join(texts.values()), count)
    return jsonify({"quiz": questions, "total": len(questions)})

# ── Stats ─────────────────────────────────────────────────────────────────────
@app.route("/stats", methods=["POST"])
def stats():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400
    per_file = {f: _get_stats(t, f) for f, t in texts.items()}
    totals   = {
        "total_words":     sum(s["words"] for s in per_file.values()),
        "total_sentences": sum(s["sentences"] for s in per_file.values()),
        "total_pages_est": sum(s["estimated_pages"] for s in per_file.values()),
        "total_read_time": sum(s["read_time_min"] for s in per_file.values()),
    }
    return jsonify({"per_file": per_file, "totals": totals})

# ── Compare ───────────────────────────────────────────────────────────────────
@app.route("/compare", methods=["POST"])
def compare():
    data = request.json or {}
    f1, f2 = data.get("file1"), data.get("file2")
    if not f1 or not f2:
        return jsonify({"error": "file1 aur file2 dono dene honge"}), 400
    texts = _load_texts([f1, f2])
    if len(texts) < 2:
        return jsonify({"error": "Dono files nahi mili"}), 400
    t1, t2 = list(texts.values())
    w1 = set(re.findall(r'[a-zA-Z]{3,}', t1.lower())) - STOP_WORDS
    w2 = set(re.findall(r'[a-zA-Z]{3,}', t2.lower())) - STOP_WORDS
    jaccard = round(len(w1 & w2) / len(w1 | w2) * 100, 1) if (w1 | w2) else 0
    common_keywords = list(w1 & w2)[:30]
    similar_pairs = []
    if ML_OK:
        try:
            s1 = split_sentences(clean_text(t1))
            s2 = split_sentences(clean_text(t2))
            vec = TfidfVectorizer(stop_words="english")
            vec.fit(s1 + s2)
            m1 = vec.transform(s1); m2 = vec.transform(s2)
            sim = cosine_similarity(m1, m2)
            for i in range(min(len(s1), 10)):
                j = sim[i].argmax()
                if sim[i][j] > 0.5:
                    similar_pairs.append({
                        "from_file1": s1[i], "from_file2": s2[j],
                        "similarity": round(float(sim[i][j]), 2),
                    })
        except: pass
    return jsonify({
        "file1": f1, "file2": f2,
        "word_overlap_pct": jaccard,
        "common_keywords": common_keywords,
        "similar_sentence_pairs": similar_pairs[:5],
        "unique_to_file1": len(w1 - w2),
        "unique_to_file2": len(w2 - w1),
    })

# ── Export DOCX ───────────────────────────────────────────────────────────────
@app.route("/export-docx", methods=["POST"])
def export_docx_route():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400
    lines   = smart_merge(texts)
    sources = list(texts.keys())
    outpath = os.path.join(OUTPUT_DIR, "studysync_notes.docx")
    if not export_docx(lines, sources, outpath):
        return jsonify({"error": "pip install python-docx"}), 500
    return jsonify({"filename": "studysync_notes.docx", "lines": len(lines)})

# ── Outputs list ──────────────────────────────────────────────────────────────
@app.route("/outputs")
def list_outputs():
    files = [{"name": f,
               "size_kb": round(os.path.getsize(os.path.join(OUTPUT_DIR, f)) / 1024, 1)}
             for f in os.listdir(OUTPUT_DIR)
             if os.path.isfile(os.path.join(OUTPUT_DIR, f))]
    return jsonify({"files": files})

# ── Download ──────────────────────────────────────────────────────────────────
@app.route("/download/<filename>")
def download(filename):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        return jsonify({"error": "File nahi mili"}), 404
    return send_file(path, as_attachment=True)

# ── Clear all ─────────────────────────────────────────────────────────────────
@app.route("/clear", methods=["POST"])
def clear():
    removed = {"uploads": [], "outputs": []}
    for f in os.listdir(UPLOAD_DIR):
        os.remove(os.path.join(UPLOAD_DIR, f)); removed["uploads"].append(f)
    for f in os.listdir(OUTPUT_DIR):
        os.remove(os.path.join(OUTPUT_DIR, f)); removed["outputs"].append(f)
    return jsonify({"cleared": removed,
                    "total": len(removed["uploads"]) + len(removed["outputs"])})

# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=" * 55)
    print("  StudySync Backend  v5")
    print(f"  URL        : http://localhost:5000")
    print(f"  Health     : http://localhost:5000/health")
    print(f"  pdfplumber : {'✓ OK' if PDF_OK  else '✗ pip install pdfplumber'}")
    print(f"  sklearn    : {'✓ OK' if ML_OK   else '✗ pip install scikit-learn'}")
    print(f"  reportlab  : {'✓ OK' if RL_OK   else '✗ pip install reportlab'}")
    print(f"  python-docx: {'✓ OK' if DOCX_OK else '✗ pip install python-docx'}")
    print(f"  Dedup      : Exact + Word Overlap ({int(WORD_OVERLAP_MIN*100)}%) + TF-IDF ({TFIDF_THRESHOLD})")
    print("=" * 55)
    app.run(debug=True, port=5000)
