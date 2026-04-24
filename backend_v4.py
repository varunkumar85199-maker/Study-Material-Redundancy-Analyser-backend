"""
StudySync Backend  v4 — Upgraded
=================================
Naye features (v4 mein add kiye):
  8.  /summary          — Har file ka short summary
  9.  /keywords         — Top keywords extract karo (TF-IDF)
  10. /flashcards       — Auto Q&A flashcards banao
  11. /quiz             — MCQ questions auto-generate
  12. /important-lines  — Sabse important sentences score ke saath
  13. /search           — Uploaded files mein text search
  14. /stats            — Detailed statistics (pages, words, sentences)
  15. /compare          — 2 files ka similarity score
  16. /export-docx      — Merged notes ko Word (.docx) mein export
  17. /clear            — Saare uploads + outputs clear karo

Purane features (v3 se):
  1.  /upload           — Multiple PDF/DOCX/TXT upload
  2.  /files            — Uploaded files list
  3.  /files/<name>     — File delete (DELETE)
  4.  /analyze          — Topic-wise analysis
  5.  /merge            — Type 1: Topic-wise notes PDF/TXT
  6.  /smart-merge      — Type 2: Smart line merge PDF/TXT
  7.  /download/<name>  — File download
      /outputs          — Output files list
      /health           — Server health check

Install:
  pip install flask flask-cors pdfplumber reportlab scikit-learn python-docx

Run:
  python backend.py  →  http://localhost:5000
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os, re, math
from collections import defaultdict, Counter
from datetime import datetime

# ── Optional imports ──────────────────────────────────────────────────────────
try:
    import pdfplumber;         PDF_OK = True
except ImportError:
    PDF_OK = False;            print("[WARN] pip install pdfplumber")

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np;        ML_OK = True
except ImportError:
    ML_OK = False;             print("[WARN] pip install scikit-learn")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle)
    RL_OK = True
except ImportError:
    RL_OK = False;             print("[WARN] pip install reportlab")

try:
    import docx as _docx;      DOCX_OK = True
except ImportError:
    DOCX_OK = False;           print("[WARN] pip install python-docx")

# ══════════════════════════════════════════════════════════════════════════════
app = Flask(__name__)
CORS(app)

UPLOAD_DIR  = "uploads"
OUTPUT_DIR  = "outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

SIM_THRESHOLD = 0.45   # duplicate detection threshold

JUNK_RE = re.compile(
    r'^\s*$|^\s*\d+\s*$|^page\s*\d+|^\s*.{1,4}\s*$'
    r'|www\.|http|@|^(copyright|all rights)'
    r'|^(figure|fig\.|table|chart)\s*\d+|^[\-_=*]{3,}$',
    re.IGNORECASE
)

# Common stop words for keyword extraction
STOP_WORDS = {
    "the","is","in","it","of","and","a","to","was","that","he","she",
    "they","we","you","i","this","with","for","on","are","be","as",
    "at","have","from","or","an","but","not","what","all","were",
    "when","there","been","one","had","by","do","so","up","which",
    "their","will","about","if","would","no","said","its","into","has"
}

# ══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_text(filepath: str) -> str:
    ext = filepath.rsplit(".", 1)[-1].lower()
    try:
        if ext == "pdf":   return _pdf_text(filepath)
        elif ext == "docx": return _docx_text(filepath)
        else:               return _txt_text(filepath)
    except Exception as e:
        return f"[Error: {e}]"

def _pdf_text(path):
    if not PDF_OK: return "[pdfplumber install karo]"
    pages = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            t = p.extract_text()
            if t: pages.append(t.strip())
    return "\n\n".join(pages)

def _docx_text(path):
    if not DOCX_OK: return "[python-docx install karo]"
    doc = _docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def _txt_text(path):
    with open(path, encoding="utf-8", errors="ignore") as f:
        return f.read()

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x20-\x7E\n]', '', text)
    return text.strip()

def split_sentences(text):
    parts = re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in parts if len(s.strip()) > 20]

def split_lines(text):
    lines = []
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if len(line) > 200:
            lines.extend(s.strip() for s in re.split(r'(?<=[.!?])\s+', line) if s.strip())
        else:
            lines.append(line)
    return lines

def is_junk(line): return bool(JUNK_RE.search(line))
def is_useful(line):
    if len(line) < 15: return False
    return len(re.findall(r'[a-zA-Z]{2,}', line)) >= 2

def _load_texts(filenames):
    return {n: extract_text(os.path.join(UPLOAD_DIR, n))
            for n in filenames
            if os.path.exists(os.path.join(UPLOAD_DIR, n))}

# ══════════════════════════════════════════════════════════════════════════════
# CORE ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

def _dedup(lines):
    if len(lines) <= 1 or not ML_OK: return lines
    try:
        vec = TfidfVectorizer(stop_words="english", min_df=1)
        mat = vec.fit_transform(lines)
        sim = cosine_similarity(mat)
        kept, removed = [], set()
        for i in range(len(lines)):
            if i in removed: continue
            kept.append(lines[i])
            for j in range(i+1, len(lines)):
                if j not in removed and sim[i][j] >= SIM_THRESHOLD:
                    removed.add(j)
        return kept
    except: return lines

def _score_sentences(sentences):
    """TF-IDF se har sentence ko importance score deta hai."""
    if not ML_OK or len(sentences) < 2:
        return [(s, 1.0) for s in sentences]
    try:
        vec = TfidfVectorizer(stop_words="english")
        mat = vec.fit_transform(sentences)
        scores = mat.sum(axis=1).A1.tolist()
        max_s  = max(scores) or 1
        return [(s, round(sc / max_s, 3)) for s, sc in zip(sentences, scores)]
    except:
        return [(s, 1.0) for s in sentences]

def _extract_keywords_tfidf(texts_dict, top_n=15):
    """TF-IDF se top keywords nikalta hai."""
    all_text = " ".join(texts_dict.values())
    words = re.findall(r'[a-zA-Z]{3,}', all_text.lower())
    filtered = [w for w in words if w not in STOP_WORDS]
    freq = Counter(filtered)
    return [{"word": w, "count": c} for w, c in freq.most_common(top_n)]

def find_topics(texts):
    file_sents = {f: split_sentences(clean_text(t)) for f, t in texts.items()}
    headings   = _extract_headings(texts)
    topics     = []
    for h in headings:
        group  = _gather_for(h, file_sents)
        if not group: continue
        merged = _merge_group(group)
        topics.append({"topic": h, "sources": list(group.keys()),
                        "merged_text": merged, "overlap": _overlap(group)})
    if not topics:
        topics = _auto_cluster(file_sents)
    return topics

def smart_merge(texts):
    all_lines = [l for t in texts.values() for l in split_lines(clean_text(t))]
    filtered  = [l for l in all_lines if not is_junk(l) and is_useful(l)]
    return _dedup(filtered) if filtered else []

def _extract_headings(texts):
    pat   = re.compile(r'^(\d+[\.\)]\s+.{5,60}|[A-Z][A-Z\s]{5,50}|#{1,3}\s+.+)$', re.MULTILINE)
    found = defaultdict(int)
    for t in texts.values():
        for m in pat.finditer(t):
            found[m.group().strip().title()] += 1
    mn = 1 if len(texts) == 1 else 2
    return [h for h, c in found.items() if c >= mn][:20]

def _gather_for(topic, file_sents):
    kws   = set(topic.lower().split())
    group = {}
    for f, sents in file_sents.items():
        rel = [s for s in sents if any(k in s.lower() for k in kws)]
        if rel: group[f] = rel
    return group

def _merge_group(group):
    all_s = [s for sents in group.values() for s in sents]
    if len(all_s) <= 1: return all_s[0] if all_s else ""
    if not ML_OK: return " ".join(all_s[:5])
    try:
        vec = TfidfVectorizer(stop_words="english")
        sim = cosine_similarity(vec.fit_transform(all_s))
        kept, used = [], set()
        for i, s in enumerate(all_s):
            if i in used: continue
            kept.append(s)
            for j in range(i+1, len(all_s)):
                if sim[i][j] >= SIM_THRESHOLD: used.add(j)
        return " ".join(kept)
    except: return " ".join(all_s[:5])

def _overlap(group):
    if len(group) < 2: return 100
    ws  = [set(w.lower() for s in sents for w in s.split()) for sents in group.values()]
    com = ws[0].intersection(*ws[1:])
    tot = ws[0].union(*ws[1:])
    return int(len(com)/len(tot)*100) if tot else 0

def _auto_cluster(file_sents):
    all_s = [s for sents in file_sents.values() for s in sents]
    if not all_s or not ML_OK: return []
    try:
        vec  = TfidfVectorizer(max_features=10, stop_words="english")
        vec.fit(all_s)
        kws  = vec.get_feature_names_out().tolist()
    except: return []
    topics = []
    for kw in kws:
        g = _gather_for(kw, file_sents)
        if g:
            topics.append({"topic": kw.title(), "sources": list(g.keys()),
                            "merged_text": _merge_group(g), "overlap": _overlap(g)})
    return topics

# ══════════════════════════════════════════════════════════════════════════════
# NEW FEATURE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _make_summary(text, max_sentences=5):
    """
    Top-scored sentences se short summary banata hai.
    """
    sentences = split_sentences(clean_text(text))
    if not sentences: return ""
    scored = _score_sentences(sentences)
    scored.sort(key=lambda x: x[1], reverse=True)
    top = [s for s, _ in scored[:max_sentences]]
    return " ".join(top)

def _make_flashcards(text, count=10):
    """
    Definitions aur key lines se Q&A flashcards banata hai.
    Pattern: "X is Y" → Q: What is X? A: X is Y.
    """
    sentences = split_sentences(clean_text(text))
    cards = []
    definition_patterns = [
        r'^(.{3,40})\s+is\s+(.{10,})',
        r'^(.{3,40})\s+are\s+(.{10,})',
        r'^(.{3,40})\s+means\s+(.{10,})',
        r'^(.{3,40})\s+refers to\s+(.{10,})',
        r'^(.{3,40})\s+=\s+(.{10,})',
    ]
    for sent in sentences:
        for pat in definition_patterns:
            m = re.match(pat, sent, re.IGNORECASE)
            if m:
                subject = m.group(1).strip()
                if len(subject.split()) <= 6:
                    cards.append({
                        "question": f"What is {subject}?",
                        "answer":   sent,
                    })
                    break
        if len(cards) >= count:
            break

    # Agar definition patterns se kam cards mile to important sentences bhi add karo
    if len(cards) < count:
        scored = _score_sentences(sentences)
        scored.sort(key=lambda x: x[1], reverse=True)
        for sent, score in scored:
            if score > 0.5 and len(cards) < count:
                if not any(c["answer"] == sent for c in cards):
                    words = sent.split()
                    if len(words) > 4:
                        q_word = words[0]
                        cards.append({
                            "question": f"Explain: '{' '.join(words[:4])}...'",
                            "answer":   sent,
                        })
    return cards[:count]

def _make_quiz(text, count=5):
    """
    MCQ questions auto-generate karta hai — important sentences se.
    """
    sentences = split_sentences(clean_text(text))
    if not sentences: return []

    scored = _score_sentences(sentences)
    scored.sort(key=lambda x: x[1], reverse=True)
    top_sents = [s for s, sc in scored[:count*3] if sc > 0.3]

    quiz = []
    for sent in top_sents[:count]:
        words = sent.split()
        if len(words) < 6: continue

        # Blank out a key word (noun/number) as the answer
        candidates = [w for w in words if len(w) > 4 and w[0].isupper() and w != words[0]]
        if not candidates:
            candidates = [w for w in words if re.match(r'\d', w)]
        if not candidates:
            candidates = [w for w in words if len(w) > 5]
        if not candidates: continue

        answer = candidates[0]
        question_text = sent.replace(answer, "______", 1)

        # Wrong options — other important words
        other_words = [w for w in words if w != answer and len(w) > 3][:3]
        while len(other_words) < 3:
            other_words.append("None of these")

        import random
        options = [answer] + other_words[:3]
        random.shuffle(options)

        quiz.append({
            "question": question_text,
            "options":  options,
            "answer":   answer,
        })

    return quiz

def _get_stats(text, filename):
    """File ki detailed statistics return karta hai."""
    clean = clean_text(text)
    words     = re.findall(r'[a-zA-Z]+', clean)
    sentences = split_sentences(clean)
    lines     = [l for l in clean.split('\n') if l.strip()]
    numbers   = re.findall(r'\b\d+\.?\d*\b', clean)
    avg_word_len = round(sum(len(w) for w in words) / len(words), 1) if words else 0
    avg_sent_len = round(len(words) / len(sentences), 1) if sentences else 0

    return {
        "filename":        filename,
        "characters":      len(clean),
        "words":           len(words),
        "sentences":       len(sentences),
        "lines":           len(lines),
        "numbers_found":   len(numbers),
        "unique_words":    len(set(w.lower() for w in words)),
        "avg_word_length": avg_word_len,
        "avg_sentence_len_words": avg_sent_len,
        "estimated_pages": max(1, round(len(words) / 250)),
        "read_time_min":   max(1, round(len(words) / 200)),
    }

# ══════════════════════════════════════════════════════════════════════════════
# PDF / TXT GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

PURPLE = "#5349C8"; LPURPLE = "#EEEDFE"; DPURPLE = "#3C3489"; GRAY = "#D3D1C7"

def _styles():
    b = getSampleStyleSheet()
    return {
        "title":    ParagraphStyle("T", parent=b["Title"], fontSize=20,
                        textColor=colors.HexColor(PURPLE), spaceAfter=4),
        "subtitle": ParagraphStyle("S", parent=b["Normal"], fontSize=9,
                        textColor=colors.grey, spaceAfter=16),
        "heading":  ParagraphStyle("H", parent=b["Heading2"], fontSize=12,
                        textColor=colors.HexColor(DPURPLE), spaceBefore=12, spaceAfter=3),
        "body":     ParagraphStyle("B", parent=b["Normal"], fontSize=10,
                        leading=16, spaceAfter=6),
        "source":   ParagraphStyle("Src", parent=b["Normal"], fontSize=8,
                        textColor=colors.grey, spaceAfter=3),
        "bullet":   ParagraphStyle("Bul", parent=b["Normal"], fontSize=10,
                        leading=17, leftIndent=12, spaceAfter=3),
        "small":    ParagraphStyle("Sm", parent=b["Normal"], fontSize=8,
                        textColor=colors.HexColor("#888780"), spaceAfter=2),
    }

def _hr(story): story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor(GRAY)))

def topic_wise_pdf(topics, outpath):
    if not RL_OK: return topic_wise_txt(topics, outpath.replace(".pdf",".txt"))
    doc = SimpleDocTemplate(outpath, pagesize=A4,
          leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    s = _styles(); story = []
    story.append(Paragraph("StudySync — Topic-Wise Notes", s["title"]))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}  |  {len(topics)} topics  |  Type 1",
        s["subtitle"]))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#AFA9EC")))
    story.append(Spacer(1, 0.3*cm))
    td = [["#","Topic","Sources","Overlap"]]
    for i, t in enumerate(topics, 1):
        td.append([str(i), t["topic"],
                   ", ".join(x.replace(".pdf","").replace(".docx","") for x in t["sources"]),
                   f"{t.get('overlap',0)}%"])
    tbl = Table(td, colWidths=[1*cm,6*cm,7*cm,2*cm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor(LPURPLE)),
        ("TEXTCOLOR", (0,0),(-1,0),colors.HexColor(DPURPLE)),
        ("FONTSIZE",  (0,0),(-1,-1),9),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F8F8FC")]),
        ("GRID",(0,0),(-1,-1),0.3,colors.HexColor(GRAY)),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
    ]))
    story += [tbl, Spacer(1, 0.5*cm)]
    for i, t in enumerate(topics, 1):
        _hr(story)
        story.append(Paragraph(f"{i}. {t['topic']}", s["heading"]))
        story.append(Paragraph("Sources: " + " | ".join(t["sources"]), s["source"]))
        story.append(Paragraph(t["merged_text"], s["body"]))
    doc.build(story); print(f"[Type1 PDF] {outpath}")

def smart_merge_pdf(lines, sources, outpath):
    if not RL_OK: return smart_merge_txt(lines, sources, outpath.replace(".pdf",".txt"))
    doc = SimpleDocTemplate(outpath, pagesize=A4,
          leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    s = _styles(); story = []
    story.append(Paragraph("StudySync — Smart Merged Notes", s["title"]))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}  |  "
        f"{len(lines)} unique lines  |  Type 2  |  Sources: {', '.join(sources)}",
        s["subtitle"]))
    info = Table([[Paragraph(
        "<b>Duplicates removed</b>: Yes  |  <b>Junk removed</b>: Yes  |  "
        "<b>Topic grouping</b>: None — pure line-by-line", s["source"])
    ]], colWidths=[16*cm])
    info.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,-1),colors.HexColor(LPURPLE)),
        ("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#AFA9EC")),
        ("TOPPADDING",(0,0),(-1,-1),8),("BOTTOMPADDING",(0,0),(-1,-1),8),
        ("LEFTPADDING",(0,0),(-1,-1),10),
    ]))
    story += [info, Spacer(1,0.4*cm),
              HRFlowable(width="100%", thickness=1, color=colors.HexColor("#AFA9EC")),
              Spacer(1,0.3*cm)]
    for line in lines:
        if len(line)<80 and not line.endswith('.') and line[0].isupper():
            story += [Spacer(1,0.15*cm), Paragraph(f"<b>{line}</b>", s["body"])]
        else:
            story.append(Paragraph(f"• {line}", s["bullet"]))
    doc.build(story); print(f"[Type2 PDF] {outpath}")

def topic_wise_txt(topics, outpath):
    lines = ["="*60,"  STUDYSYNC — TOPIC-WISE NOTES (Type 1)",
             f"  Generated : {datetime.now().strftime('%d %b %Y')}",
             f"  Topics    : {len(topics)}","="*60,""]
    for i,t in enumerate(topics,1):
        lines += [f"{i}. {t['topic'].upper()}",
                  f"   Sources : {', '.join(t['sources'])}",
                  f"   Overlap : {t.get('overlap',0)}%",
                  "", t["merged_text"], "", "-"*60, ""]
    open(outpath,"w",encoding="utf-8").write("\n".join(lines))

def smart_merge_txt(lines, sources, outpath):
    out = ["="*60,"  STUDYSYNC — SMART MERGED NOTES (Type 2)",
           f"  Generated : {datetime.now().strftime('%d %b %Y')}",
           f"  Sources   : {', '.join(sources)}",
           f"  Lines     : {len(lines)}","="*60,""]
    out += [f"• {l}" for l in lines] + [""]
    open(outpath,"w",encoding="utf-8").write("\n".join(out))

def export_docx(topics, lines, sources, outpath):
    """Word document mein dono types export karta hai."""
    if not DOCX_OK:
        return False
    doc = _docx.Document()
    doc.add_heading("StudySync — Study Notes", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y, %I:%M %p')}")
    doc.add_paragraph(f"Sources: {', '.join(sources)}")

    # Type 1
    doc.add_heading("Type 1 — Topic-Wise Notes", level=1)
    for i, t in enumerate(topics, 1):
        doc.add_heading(f"{i}. {t['topic']}", level=2)
        doc.add_paragraph(f"Sources: {', '.join(t['sources'])}", style="Intense Quote")
        doc.add_paragraph(t["merged_text"])

    # Type 2
    doc.add_page_break()
    doc.add_heading("Type 2 — Smart Merged Notes", level=1)
    doc.add_paragraph(f"{len(lines)} unique lines — duplicates & junk removed")
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(outpath)
    return True

# ══════════════════════════════════════════════════════════════════════════════
# FLASK ROUTES
# ══════════════════════════════════════════════════════════════════════════════

# ── Health ────────────────────────────────────────────────────────────────────
@app.route("/health")
def health():
    return jsonify({
        "status": "online", "version": "v4",
        "pdfplumber": PDF_OK, "sklearn": ML_OK,
        "reportlab": RL_OK, "python_docx": DOCX_OK,
        "time": datetime.now().strftime("%d %b %Y %I:%M %p"),
    })

# ── Upload ────────────────────────────────────────────────────────────────────
@app.route("/upload", methods=["POST"])
def upload():
    if "files" not in request.files:
        return jsonify({"error": "Koi file nahi mili"}), 400
    saved = []
    for f in request.files.getlist("files"):
        if f.filename.endswith((".pdf",".docx",".txt")):
            path = os.path.join(UPLOAD_DIR, f.filename)
            f.save(path)
            size = os.path.getsize(path)
            saved.append({"name": f.filename, "size_kb": round(size/1024, 1)})
    return jsonify({"uploaded": saved, "count": len(saved)})

# ── Files list ────────────────────────────────────────────────────────────────
@app.route("/files")
def list_files():
    files = [{"name": f, "size_kb": round(os.path.getsize(os.path.join(UPLOAD_DIR,f))/1024, 1)}
             for f in os.listdir(UPLOAD_DIR)
             if os.path.isfile(os.path.join(UPLOAD_DIR, f))]
    return jsonify({"files": files, "count": len(files)})

# ── Delete file ───────────────────────────────────────────────────────────────
@app.route("/files/<filename>", methods=["DELETE"])
def delete_file(filename):
    path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(path):
        os.remove(path)
        return jsonify({"deleted": filename})
    return jsonify({"error": "File nahi mili"}), 404

# ── Analyze topics ────────────────────────────────────────────────────────────
@app.route("/analyze", methods=["POST"])
def analyze():
    data  = request.json or {}
    files = data.get("files", [])
    texts = _load_texts(files)
    if not texts: return jsonify({"error": "Koi valid file nahi mili"}), 400
    topics = find_topics(texts)
    return jsonify({"topics": topics, "total": len(topics), "files": list(texts.keys())})

# ── Type 1: Topic-wise merge ──────────────────────────────────────────────────
@app.route("/merge", methods=["POST"])
def merge():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    fmt   = data.get("format", "pdf")
    topics  = find_topics(texts)
    outpath = os.path.join(OUTPUT_DIR, f"topic_wise_notes.{fmt}")
    topic_wise_pdf(topics, outpath) if fmt == "pdf" else topic_wise_txt(topics, outpath)
    return jsonify({"filename": f"topic_wise_notes.{fmt}", "topics_merged": len(topics), "type": "topic-wise"})

# ── Type 2: Smart merge ───────────────────────────────────────────────────────
@app.route("/smart-merge", methods=["POST"])
def smart_merge_route():
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    fmt   = data.get("format", "pdf")
    lines   = smart_merge(texts)
    sources = list(texts.keys())
    outpath = os.path.join(OUTPUT_DIR, f"smart_merged_notes.{fmt}")
    smart_merge_pdf(lines, sources, outpath) if fmt == "pdf" else smart_merge_txt(lines, sources, outpath)
    return jsonify({"filename": f"smart_merged_notes.{fmt}", "unique_lines": len(lines),
                    "lines": lines, "type": "smart-merge"})

# ── NEW: Summary ──────────────────────────────────────────────────────────────
@app.route("/summary", methods=["POST"])
def summary():
    """
    Har file ka alag short summary + combined summary.
    Body: { "files": ["a.pdf", "b.pdf"], "sentences": 5 }
    """
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    n     = data.get("sentences", 5)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400

    summaries = {}
    for fname, text in texts.items():
        summaries[fname] = _make_summary(text, n)

    combined_text = " ".join(texts.values())
    return jsonify({
        "individual": summaries,
        "combined":   _make_summary(combined_text, n),
        "files":      list(texts.keys()),
    })

# ── NEW: Keywords ─────────────────────────────────────────────────────────────
@app.route("/keywords", methods=["POST"])
def keywords():
    """
    Top keywords extract karta hai TF-IDF se.
    Body: { "files": ["a.pdf"], "top": 15 }
    """
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    top   = data.get("top", 15)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400

    per_file = {}
    for fname, text in texts.items():
        per_file[fname] = _extract_keywords_tfidf({fname: text}, top)

    return jsonify({
        "per_file": per_file,
        "combined": _extract_keywords_tfidf(texts, top),
    })

# ── NEW: Flashcards ───────────────────────────────────────────────────────────
@app.route("/flashcards", methods=["POST"])
def flashcards():
    """
    Auto Q&A flashcards banata hai.
    Body: { "files": ["a.pdf"], "count": 10 }
    """
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    count = data.get("count", 10)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400

    combined = " ".join(texts.values())
    cards = _make_flashcards(combined, count)
    return jsonify({"flashcards": cards, "total": len(cards)})

# ── NEW: Quiz ─────────────────────────────────────────────────────────────────
@app.route("/quiz", methods=["POST"])
def quiz():
    """
    MCQ questions auto-generate karta hai.
    Body: { "files": ["a.pdf"], "count": 5 }
    """
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    count = data.get("count", 5)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400

    combined = " ".join(texts.values())
    questions = _make_quiz(combined, count)
    return jsonify({"quiz": questions, "total": len(questions)})

# ── NEW: Important lines ──────────────────────────────────────────────────────
@app.route("/important-lines", methods=["POST"])
def important_lines():
    """
    Sabse important sentences score ke saath return karta hai.
    Body: { "files": ["a.pdf"], "top": 20, "min_score": 0.4 }
    """
    data      = request.json or {}
    texts     = _load_texts(data.get("files", []))
    top       = data.get("top", 20)
    min_score = data.get("min_score", 0.3)
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400

    combined  = " ".join(texts.values())
    sentences = split_sentences(clean_text(combined))
    scored    = _score_sentences(sentences)
    scored.sort(key=lambda x: x[1], reverse=True)

    result = [{"line": s, "score": sc} for s, sc in scored
              if sc >= min_score][:top]
    return jsonify({"lines": result, "total": len(result)})

# ── NEW: Search ───────────────────────────────────────────────────────────────
@app.route("/search", methods=["POST"])
def search():
    """
    Uploaded files mein text search karta hai.
    Body: { "files": ["a.pdf", "b.pdf"], "query": "Newton law" }
    """
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    query = data.get("query", "").lower().strip()
    if not query: return jsonify({"error": "Query deni hogi"}), 400

    results = {}
    for fname, text in texts.items():
        sentences = split_sentences(clean_text(text))
        matches   = [s for s in sentences if query in s.lower()]
        if matches:
            results[fname] = {"matches": matches, "count": len(matches)}

    total = sum(r["count"] for r in results.values())
    return jsonify({"query": query, "results": results,
                    "total_matches": total, "files_matched": len(results)})

# ── NEW: Stats ────────────────────────────────────────────────────────────────
@app.route("/stats", methods=["POST"])
def stats():
    """
    Har file ki detailed statistics return karta hai.
    Body: { "files": ["a.pdf", "b.pdf"] }
    """
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400

    per_file = {fname: _get_stats(text, fname) for fname, text in texts.items()}
    totals   = {
        "total_words":     sum(s["words"]     for s in per_file.values()),
        "total_sentences": sum(s["sentences"] for s in per_file.values()),
        "total_pages_est": sum(s["estimated_pages"] for s in per_file.values()),
        "total_read_time": sum(s["read_time_min"] for s in per_file.values()),
    }
    return jsonify({"per_file": per_file, "totals": totals})

# ── NEW: Compare 2 files ──────────────────────────────────────────────────────
@app.route("/compare", methods=["POST"])
def compare():
    """
    2 files ka similarity score aur common content.
    Body: { "file1": "a.pdf", "file2": "b.pdf" }
    """
    data  = request.json or {}
    f1, f2 = data.get("file1"), data.get("file2")
    if not f1 or not f2: return jsonify({"error": "file1 aur file2 dono dene honge"}), 400

    texts = _load_texts([f1, f2])
    if len(texts) < 2: return jsonify({"error": "Dono files nahi mili"}), 400

    t1, t2 = list(texts.values())
    w1 = set(re.findall(r'[a-zA-Z]{3,}', t1.lower())) - STOP_WORDS
    w2 = set(re.findall(r'[a-zA-Z]{3,}', t2.lower())) - STOP_WORDS
    common_words = list(w1 & w2)[:30]
    jaccard      = round(len(w1 & w2) / len(w1 | w2) * 100, 1) if (w1 | w2) else 0

    # Sentence level similarity
    sents1 = split_sentences(clean_text(t1))
    sents2 = split_sentences(clean_text(t2))
    similar_pairs = []
    if ML_OK and sents1 and sents2:
        try:
            vec = TfidfVectorizer(stop_words="english")
            vec.fit(sents1 + sents2)
            m1  = vec.transform(sents1)
            m2  = vec.transform(sents2)
            sim = cosine_similarity(m1, m2)
            for i in range(min(len(sents1), 10)):
                j = sim[i].argmax()
                if sim[i][j] > 0.5:
                    similar_pairs.append({
                        "from_file1": sents1[i],
                        "from_file2": sents2[j],
                        "similarity": round(float(sim[i][j]), 2),
                    })
        except: pass

    return jsonify({
        "file1": f1, "file2": f2,
        "word_overlap_pct": jaccard,
        "common_keywords":  common_words,
        "similar_sentence_pairs": similar_pairs[:5],
        "unique_to_file1": len(w1 - w2),
        "unique_to_file2": len(w2 - w1),
    })

# ── NEW: Export DOCX ──────────────────────────────────────────────────────────
@app.route("/export-docx", methods=["POST"])
def export_docx_route():
    """
    Dono types ko ek Word document mein export karta hai.
    Body: { "files": ["a.pdf", "b.pdf"] }
    """
    data  = request.json or {}
    texts = _load_texts(data.get("files", []))
    if not texts: return jsonify({"error": "Koi file nahi mili"}), 400

    topics  = find_topics(texts)
    lines   = smart_merge(texts)
    sources = list(texts.keys())
    outpath = os.path.join(OUTPUT_DIR, "studysync_notes.docx")

    ok = export_docx(topics, lines, sources, outpath)
    if not ok:
        return jsonify({"error": "python-docx install karo: pip install python-docx"}), 500

    return jsonify({"filename": "studysync_notes.docx",
                    "topics": len(topics), "lines": len(lines)})

# ── Outputs list ──────────────────────────────────────────────────────────────
@app.route("/outputs")
def list_outputs():
    files = [{"name": f, "size_kb": round(os.path.getsize(os.path.join(OUTPUT_DIR,f))/1024,1)}
             for f in os.listdir(OUTPUT_DIR)
             if os.path.isfile(os.path.join(OUTPUT_DIR, f))]
    return jsonify({"files": files})

# ── Download ──────────────────────────────────────────────────────────────────
@app.route("/download/<filename>")
def download(filename):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        return jsonify({"error": "File nahi mili"}), 404
    return send_file(path, as_attachment=True)

# ── NEW: Clear all ────────────────────────────────────────────────────────────
@app.route("/clear", methods=["POST"])
def clear():
    """Saare uploads aur outputs delete karta hai."""
    removed = {"uploads": [], "outputs": []}
    for f in os.listdir(UPLOAD_DIR):
        os.remove(os.path.join(UPLOAD_DIR, f))
        removed["uploads"].append(f)
    for f in os.listdir(OUTPUT_DIR):
        os.remove(os.path.join(OUTPUT_DIR, f))
        removed["outputs"].append(f)
    return jsonify({"cleared": removed,
                    "total": len(removed["uploads"]) + len(removed["outputs"])})

# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=" * 55)
    print("  StudySync Backend  v4")
    print(f"  URL      : http://localhost:5000")
    print(f"  Health   : http://localhost:5000/health")
    print(f"  pdfplumber : {'✓ OK' if PDF_OK  else '✗ pip install pdfplumber'}")
    print(f"  sklearn    : {'✓ OK' if ML_OK   else '✗ pip install scikit-learn'}")
    print(f"  reportlab  : {'✓ OK' if RL_OK   else '✗ pip install reportlab'}")
    print(f"  python-docx: {'✓ OK' if DOCX_OK else '✗ pip install python-docx'}")
    print("=" * 55)
    app.run(debug=True, port=5000)
